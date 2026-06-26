#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
根据PDF内容填写在职证明DOC
修改项：护照号→ER606835，公司地址→海南省海口市秀英区秀英街道港澳大道7号
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

FONT_NAME = "微软雅黑"

def set_run_font(run, font_name=FONT_NAME, font_size=12, bold=False):
    """设置run的字体（含东亚字体域）"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    # 设置东亚字体
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)

def set_para_font(para, font_name=FONT_NAME):
    """设置段落默认字体"""
    pPr = para._p.get_or_add_pPr()
    rPr = pPr.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)

def add_blank_line(doc, height_pt=18):
    """添加一个空白段落作为占位"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.font.size = Pt(height_pt)
    return p

# 创建新文档
doc = Document()

# 页面边距
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(3)
section.right_margin = Cm(2)

# 全局样式设置
style = doc.styles['Normal']
style.font.name = FONT_NAME
style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
style._element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)

def add_para(doc, text="", bold=False, font_size=12, align=WD_ALIGN_PARAGRAPH.LEFT, 
             space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run_font(run, font_name=FONT_NAME, font_size=font_size, bold=bold)
    else:
        run = p.add_run()
        run.font.size = Pt(font_size)
    return p

def add_mixed_para(doc, parts, font_size=12, align=WD_ALIGN_PARAGRAPH.LEFT,
                   space_before=0, space_after=6, first_line_indent=0):
    """
    parts: [(text, bold), ...] 或 [(text, bold, underline), ...]
    """
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(first_line_indent)
    for item in parts:
        text = item[0]
        bold = item[1] if len(item) > 1 else False
        run = p.add_run(text)
        set_run_font(run, font_name=FONT_NAME, font_size=font_size, bold=bold)
    return p

# ===== 文档内容 =====

# 公司名称
add_para(doc, "华检联（海南）检测技术有限公司", bold=True, font_size=14,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=8)

# 分隔线
add_para(doc, "─────────────────────────────────────────────────────",
         font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=8)

# 标题
add_para(doc, "在  职  证  明", bold=True, font_size=16,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=20)

# 致送机构
add_mixed_para(doc, [("日本驻广州大使馆/总领事馆签证处：", False)],
               font_size=12, space_before=0, space_after=16, first_line_indent=24)

# 正文第一段
add_mixed_para(doc,
    [
        ("符发", True),
        ("（护照：", False),
        ("ER606835", True),
        ("，出生日期", False),
        ("1979年11月27日", True),
        ("），自", False),
        ("2016年10月", True),
        ("迄今在我公司工作，现任职务", False),
        ("董事长", True),
        ("，年薪为", False),
        ("25万元", True),
        ("，", False),
    ],
    font_size=12, space_before=0, space_after=6, first_line_indent=24
)

# 正文第二段
add_mixed_para(doc,
    [
        ("我公司同意其自", False),
        ("2025年01月12日", True),
        ("至", False),
        ("2025年01月22日", True),
        ("前往", False),
        ("日本", True),
        ("旅游出行，费用由其自行承担。我司保证其在此期间遵守当地法律，并保留其职务至归国。", False),
    ],
    font_size=12, space_before=0, space_after=30, first_line_indent=24
)

# 单位名称
add_mixed_para(doc, [("单位名称: ", False), ("华检联（海南）检测技术有限公司", True)],
               font_size=12, space_before=0, space_after=8, first_line_indent=24)

# 公章标注
add_para(doc, "（公司公章）", font_size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=8)

# 单位地址（已更新）
add_mixed_para(doc, [("单位地址: ", False), ("海南省海口市秀英区秀英街道港澳大道7号", True)],
               font_size=12, space_before=0, space_after=8, first_line_indent=24)

# 领导签字
add_mixed_para(doc, [("领导签字：", False), ("郑家云", True)],
               font_size=12, space_before=0, space_after=8, first_line_indent=24)

# 联系电话
add_mixed_para(doc, [("联系电话：", False), ("17384653707", True)],
               font_size=12, space_before=0, space_after=8, first_line_indent=24)

# 日期
add_mixed_para(doc, [("日期：", False)],
               font_size=12, space_before=0, space_after=6, first_line_indent=24)

# 保存
output_path = "/Users/fasimac/Downloads/在职证明_已填写.docx"
doc.save(output_path)
print(f"文件已保存: {output_path}")

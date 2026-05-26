#!/usr/bin/env python3
"""
脱敏处理华检联公司管理制度：
1. 金额数字替换为***
2. "华检联"替换为"HJL"
3. 同名文件只保留docx
"""
import os
import re
import shutil
import zipfile
import subprocess

# 配置
SOURCE_DIR = "/Users/fasimac/Desktop/管理/华检联公司管理制度"
OUTPUT_DIR = "/Users/fasimac/.qclaw/workspace/HJL管理制度_脱敏版"
ZIP_PATH = "/Users/fasimac/.qclaw/workspace/HJL管理制度_脱敏版.zip"

def desensitize_text(text):
    """对文本进行脱敏处理"""
    if not text:
        return text
    
    # 1. 先替换公司名称
    result = text.replace("华检联", "HJL").replace("华检联公司", "HJL公司")
    
    # 2. 金额数字和百分比替换 - 覆盖所有薪资相关场景
    
    # 替换模式列表（按优先级）
    replacements = [
        # 标准货币格式
        (r'¥\s*(\d+(?:\.\d+)?)', '***'),
        (r'\$\s*(\d+(?:\.\d+)?)', '***'),
        (r'(\d+(?:\.\d+)?)\s*元(?:/月|/年|/天|/人|/次|/小时|/平米)?', '***'),
        (r'人民币\s*(\d+(?:\.\d+)?)', '***'),
        (r'(\d+(?:\.\d+)?)\s*万\s*元', '***'),
        (r'(\d+(?:\.\d+)?)\s*万', '***'),
        (r'(\d+(?:\.\d+)?)\s*千', '***'),
        
        # 薪资关键词后的数字
        (r'(?:基本工资|岗位工资|绩效工资|工龄工资|学历工资|职称工资|职务工资)\s*(\d+(?:\.\d+)?)', '***'),
        (r'(?:补贴|津贴|奖金|提成|分红)\s*(\d+(?:\.\d+)?)', '***'),
        (r'(?:工资|薪资|薪酬)\s*标准\s*(\d+(?:\.\d+)?)', '***'),
        (r'(?:月薪|日薪|时薪|周薪)\s*(\d+(?:\.\d+)?)', '***'),
        
        # 含"约"或"按"的金额描述
        (r'约\s*(\d+(?:\.\d+)?)', '***'),
        (r'按\s*(\d+(?:\.\d+)?)\s*(?:元|%)', '***'),
        
        # 百分比（薪资相关）
        (r'(\d+(?:\.\d+)?)\s*%', '***'),
        
        # 独立数字行（3位以上，很可能是金额）
        (r'^\s*(\d{3,6}(?:\.\d{1,2})?)\s*$', '***'),
        
        # 数字+单位
        (r'(\d+(?:\.\d+)?)\s*(?:元/月|元/人|元/天|元/次|元/小时|元/年|元/平米|万元/年)', '***'),
        
        # 元/月、元/人等后缀在前面
        (r'(?:元/月|元/人|元/天|元/次|元/小时)\s*(\d+(?:\.\d+)?)', '***'),
    ]
    
    for pattern, replacement in replacements:
        result = re.sub(pattern, replacement, result)
    
    return result

def convert_doc_to_docx(src_path, dst_path):
    """使用textutil将doc转换为docx"""
    try:
        cmd = ['textutil', '-convert', 'docx', '-output', dst_path, src_path]
        subprocess.run(cmd, check=True, capture_output=True, timeout=60)
        return True
    except Exception as e:
        print(f"    转换失败: {e}")
        return False

def process_docx_file(src_path, dst_path):
    """处理docx文件"""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        
        doc = Document(src_path)
        modified = False
        
        # 处理所有段落
        for paragraph in doc.paragraphs:
            original = paragraph.text
            new_text = desensitize_text(original)
            if new_text != original:
                for run in paragraph.runs:
                    run.text = desensitize_text(run.text)
                modified = True
        
        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original = paragraph.text
                        new_text = desensitize_text(original)
                        if new_text != original:
                            for run in paragraph.runs:
                                run.text = desensitize_text(run.text)
                            modified = True
        
        # 设置中文字体
        if modified:
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    try:
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    except:
                        pass
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                try:
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                except:
                                    pass
        
        doc.save(dst_path)
        return True
    except Exception as e:
        print(f"  ✗ 处理失败: {e}")
        return False

def copy_and_desensitize_file(src_path, dst_path):
    """复制并脱敏文件"""
    ext = os.path.splitext(src_path)[1].lower()
    
    # doc文件先转换为docx
    if ext == '.doc':
        temp_docx = '/tmp/temp_convert_' + str(os.getpid()) + '.docx'
        if convert_doc_to_docx(src_path, temp_docx):
            success = process_docx_file(temp_docx, dst_path.replace('.doc', '.docx'))
            os.remove(temp_docx)
            if success:
                print(f"  ✓ 已处理: {os.path.basename(src_path)}")
                return 'converted'
        return 'failed'
    
    # docx文件直接处理
    if ext == '.docx':
        success = process_docx_file(src_path, dst_path)
        if success:
            print(f"  ✓ 已处理: {os.path.basename(src_path)}")
            return 'success'
        return 'failed'
    
    # PDF文件直接复制
    elif ext == '.pdf':
        shutil.copy2(src_path, dst_path)
        print(f"  ○ 跳过PDF: {os.path.basename(src_path)}")
        return 'pdf'
    
    else:
        shutil.copy2(src_path, dst_path)
        return 'success'

def process_directory():
    """处理整个目录"""
    print(f"源目录: {SOURCE_DIR}")
    print(f"输出目录: {OUTPUT_DIR}")
    print("-" * 50)
    
    # 清理旧输出
    if os.path.exists(OUTPUT_DIR):
        shutil.rmtree(OUTPUT_DIR)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    # 获取所有文件
    all_files = []
    for root, dirs, files in os.walk(SOURCE_DIR):
        dirs[:] = [d for d in dirs if not d.startswith('.')]
        for f in files:
            if not f.startswith('.') and not f.endswith('.DS_Store'):
                all_files.append(os.path.join(root, f))
    
    print(f"找到 {len(all_files)} 个文件\n")
    
    # 跟踪已处理的文件（用于去重）
    processed_basenames = {}
    
    for src_file in all_files:
        rel_path = os.path.relpath(src_file, SOURCE_DIR)
        basename = os.path.basename(rel_path)
        name_without_ext = os.path.splitext(basename)[0]
        
        ext = os.path.splitext(src_file)[1].lower()
        
        # 跳过PDF
        if ext == '.pdf':
            dst_file = os.path.join(OUTPUT_DIR, rel_path)
            os.makedirs(os.path.dirname(dst_file), exist_ok=True)
            shutil.copy2(src_file, dst_file)
            print(f"  ○ 跳过PDF: {basename}")
            continue
        
        # 确定输出路径（统一使用docx）
        if ext == '.doc':
            dst_name = name_without_ext + '.docx'
            rel_path = rel_path.replace(basename, dst_name)
        
        dst_file = os.path.join(OUTPUT_DIR, rel_path)
        os.makedirs(os.path.dirname(dst_file), exist_ok=True)
        
        # 检查是否已有同名文件
        if name_without_ext in processed_basenames:
            print(f"  ○ 跳过重复: {basename} (已有{processed_basenames[name_without_ext]})")
            continue
        
        # 处理文件
        result = copy_and_desensitize_file(src_file, dst_file)
        if result in ['success', 'converted']:
            processed_basenames[name_without_ext] = '.docx'
    
    print("-" * 50)
    print(f"处理完成: {len(processed_basenames)} 个文件（去重后）")
    return processed_basenames

def create_zip():
    """创建zip压缩包"""
    print(f"\n创建压缩包: {ZIP_PATH}")
    
    if os.path.exists(ZIP_PATH):
        os.remove(ZIP_PATH)
    
    file_count = 0
    with zipfile.ZipFile(ZIP_PATH, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(OUTPUT_DIR):
            dirs[:] = [d for d in dirs if not d.startswith('.')]
            for file in files:
                if not file.startswith('.') and file != '.DS_Store' and file.endswith('.docx'):
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, OUTPUT_DIR)
                    zipf.write(file_path, arcname)
                    file_count += 1
    
    print(f"✓ 压缩包已创建: {ZIP_PATH}")
    print(f"  文件数: {file_count} 个")
    print(f"  文件大小: {os.path.getsize(ZIP_PATH) / 1024:.1f} KB")

def verify_result():
    """验证脱敏结果"""
    print("\n" + "="*50)
    print("验证脱敏效果")
    print("="*50)
    
    from docx import Document
    
    # 检查薪资管理制度
    salary_file = f"{OUTPUT_DIR}/制度2021/202106薪资管理制度.docx"
    if os.path.exists(salary_file):
        doc = Document(salary_file)
        print(f"\n薪资管理制度检查:")
        
        # 检查纯数字段落
        digit_count = 0
        for p in doc.paragraphs:
            text = p.text.strip()
            if text.isdigit() and len(text) >= 3:
                print(f"  ✗ 残留数字: {text}")
                digit_count += 1
        
        # 检查百分比
        pct_count = 0
        for p in doc.paragraphs:
            text = p.text
            if re.search(r'\d+%', text):
                print(f"  ✗ 残留百分比: {text[:60]}")
                pct_count += 1
        
        if digit_count == 0 and pct_count == 0:
            print("  ✓ 脱敏成功！")
        else:
            print(f"  ✗ 发现 {digit_count} 个数字 + {pct_count} 个百分比残留")

if __name__ == "__main__":
    process_directory()
    create_zip()
    verify_result()
    print("\n全部完成!")
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""FIRE财务自由实践指南 Word文档生成"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# 颜色常量
TITLE = RGBColor(0x1F, 0x38, 0x64)    # 深蓝
ACCENT = RGBColor(0x2E, 0x75, 0xB6)  # 蓝
WARN = RGBColor(0xC0, 0x00, 0x00)    # 红
TEXT = RGBColor(0x2C, 0x2C, 0x2C)    # 深灰
SUB = RGBColor(0x59, 0x59, 0x59)    # 副标题灰
GREEN = RGBColor(0x2E, 0x7D, 0x32)  # 绿
GOLD = RGBColor(0x8B, 0x69, 0x14)    # 金

BG_H = "D6E4F0"     # 表头背景
BG_A = "F5F9FD"     # 交替行
WARN_BG = "FFF5F5"  # 红色背景
GREEN_BG = "F0FFF4" # 绿色背景
BLUE_BG = "EBF3FB"  # 蓝色背景
YELLOW_BG = "FFF8E1" # 黄色背景
YELLOW2_BG = "FFF3CD" # 深黄背景

def hex_to_rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def set_cell_bg(cell, hex_color):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_para_shading(para, hex_color):
    """设置段落背景色"""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def set_para_bottom_border(para, color="2E75B6", size="8"):
    """设置段落下边框"""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_run(para, text, bold=False, size=22, color=None, font="宋体"):
    run = para.add_run(text)
    run.font.name = font
    run._r.rPr.rFonts.set(qn('w:eastAsia'), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run

def new_para(doc_or_parent, align=WD_ALIGN_PARAGRAPH.LEFT, sp_before=0, sp_after=200, children=None):
    """创建段落"""
    if hasattr(doc_or_parent, 'add_paragraph'):
        p = doc_or_parent.add_paragraph()
    else:
        p = doc_or_parent._p.add_p()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sp_before / 20)
    p.paragraph_format.space_after = Pt(sp_after / 20)
    return p

def heading_para(doc, text, level=1):
    """标题段落"""
    sizes = {1: 20, 2: 16, 3: 14}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    colors = {1: TITLE, 2: ACCENT, 3: TEXT}
    sps = {1: (14, 6), 2: (10, 5), 3: (8, 4)}
    bolds = {1: True, 2: True, 3: True}

    p.paragraph_format.space_before = Pt(sps[level][0])
    p.paragraph_format.space_after = Pt(sps[level][1])

    run = p.add_run(text)
    run.font.name = "宋体"
    run._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    run.font.size = Pt(sizes[level])
    run.bold = bolds[level]
    run.font.color.rgb = colors[level]

    if level == 1:
        set_para_bottom_border(p, "2E75B6", "8")
    return p

def normal_para(doc, text, bold=False, size=22, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, sp_before=0, sp_after=6):
    """普通文本段落"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sp_before)
    p.paragraph_format.space_after = Pt(sp_after)
    run = p.add_run(text)
    run.font.name = "宋体"
    run._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def box_para(doc, text, size=28, bold=True, color=None, bg=BLUE_BG):
    """居中带背景框的段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    set_para_shading(p, bg)
    r = p.add_run(text)
    r.font.name = "宋体"
    r._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = color or TITLE
    return p

def note_para(doc, text, color=None, bg=YELLOW2_BG, size=20):
    """警示/备注段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    set_para_shading(p, bg)
    r = p.add_run(text)
    r.font.name = "宋体"
    r._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    r.font.size = Pt(size)
    r.font.color.rgb = color or GOLD
    return p

def spacer(doc, before=0, after=4):
    """空段落（间距）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p

def bullet_para(doc, text, size=22, color=None):
    """项目符号段落"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "宋体"
    r._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p

def numbered_para(doc, num, text, size=22):
    """有序列表段落"""
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "宋体"
    r._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    r.font.size = Pt(size)
    return p

def make_table(doc, headers, rows, col_widths):
    """创建表格"""
    n_cols = len(headers)
    t = doc.add_table(rows=len(rows)+1, cols=n_cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'

    # 设置列宽
    for i, w in enumerate(col_widths):
        for j, cell in enumerate(t.columns[i].cells):
            cell.width = Cm(w)

    # 表头行
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        set_cell_bg(cell, BG_H)
        set_cell_borders(cell, "BFBFBF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = "宋体"
        r._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        r.font.size = Pt(20)
        r.bold = True
        r.font.color.rgb = TITLE

    # 数据行
    for row_idx, row_data in enumerate(rows):
        bg = BG_A if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, val in enumerate(row_data):
            cell = t.rows[row_idx+1].cells[col_idx]
            cell.text = ''
            set_cell_bg(cell, bg)
            set_cell_borders(cell, "CCCCCC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.name = "宋体"
            r._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
            r.font.size = Pt(20)
    return t

def colored_table_cell(cell, text, bg="FFFFFF", bold=False, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    """彩色单元格"""
    cell.text = ''
    set_cell_bg(cell, bg)
    set_cell_borders(cell, "CCCCCC")
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    r.font.name = "宋体"
    r._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    r.font.size = Pt(20)
    r.bold = bold
    if color:
        r.font.color.rgb = color

def two_col_table(doc, rows_data, col_widths=[5, 9]):
    """两列表格"""
    t = doc.add_table(rows=len(rows_data), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for row_i, (left, right) in enumerate(rows_data):
        bg = BG_A if row_i % 2 == 0 else "FFFFFF"
        for ci, val in enumerate([left, right]):
            cell = t.rows[row_i].cells[ci]
            colored_table_cell(cell, val, bg=bg)
    return t

# ===== 生成文档 =====
doc = Document()

# 全局字体设置
style = doc.styles['Normal']
style.font.name = '宋体'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(22)

# 页面设置：A4
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)

# 页眉
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("FIRE运动 · 财务自由实践指南    2026年4月")
hr.font.name = "宋体"
hr._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
hr.font.size = Pt(9)
hr.font.color.rgb = SUB
set_para_bottom_border(hp, "2E75B6", "4")

# 页脚
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run()
fr.font.name = "宋体"
fr._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
fr.font.size = Pt(9)
fr.font.color.rgb = SUB
fr2 = fp.add_run()
fr2.font.name = "宋体"
fr2._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
fr2.font.size = Pt(9)
fr2.font.color.rgb = SUB
set_para_bottom_border(fp, "2E75B6", "4")

# ===== 封面 =====
spacer(doc, 30, 0)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
r = p.add_run("时间自由")
r.font.name = "宋体"
r._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
r.font.size = Pt(36)
r.bold = True
r.font.color.rgb = TITLE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(4)
p2.paragraph_format.space_after = Pt(12)
r2 = p2.add_run("重新定义财务自由的本质")
r2.font.name = "宋体"
r2._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
r2.font.size = Pt(18)
r2.font.color.rgb = ACCENT

spacer(doc, 4, 12)
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("内容来源：元宝AI（微信视频号解读）  |  整理加工：QClaw")
r3.font.name = "宋体"
r3._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
r3.font.size = Pt(9)
r3.font.color.rgb = SUB

# ===== 一、核心概念重构 =====
doc.add_page_break()
heading_para(doc, "一、核心概念重构", 1)

heading_para(doc, "▌ 传统误区", 2)
p_err = normal_para(doc, "将财务自由等同于巨额财富积累（如1000万/1亿）——这是最大的思维陷阱。",
                    bold=True, size=22, color=WARN)
spacer(doc, 3, 3)

heading_para(doc, "▌ 本质定义", 2)
spacer(doc, 4, 2)
box_para(doc, "被动收入 > 生活开支  =  时间自主权", size=16, bold=True, color=TITLE, bg=BLUE_BG)
spacer(doc, 4, 3)

heading_para(doc, "▌ 关键转变", 2)
spacer(doc, 3, 2)
# 从\u201c赚更多\u201d → \u201c花更少\u201d 的思维升级
p_arrow = doc.add_paragraph()
p_arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_arrow.paragraph_format.space_before = Pt(4)
p_arrow.paragraph_format.space_after = Pt(8)
r_left = p_arrow.add_run("从 ")
r_left.font.name = "宋体"
r_left._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
r_left.font.size = Pt(14)
r_left.font.color.rgb = SUB
r_warn = p_arrow.add_run('\u201c赚更多\u201d')
r_warn.font.name = "宋体"
r_warn._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
r_warn.font.size = Pt(14)
r_warn.bold = True
r_warn.font.color.rgb = WARN
r_mid = p_arrow.add_run("  →  ")
r_mid.font.name = "宋体"
r_mid._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
r_mid.font.size = Pt(14)
r_mid.font.color.rgb = SUB
r_green = p_arrow.add_run('\u201c花更少\u201d')
r_green.font.name = "宋体"
r_green._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
r_green.font.size = Pt(14)
r_green.bold = True
r_green.font.color.rgb = GREEN
r_right = p_arrow.add_run("  的思维升级")
r_right.font.name = "宋体"
r_right._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
r_right.font.size = Pt(14)
r_right.font.color.rgb = SUB

spacer(doc, 4, 4)

# 思维对比表
t2 = doc.add_table(rows=5, cols=3)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
t2.style = 'Table Grid'
headers2 = ["维度", "❌ 传统思维", "✅ FIRE思维"]
for i, h in enumerate(headers2):
    colored_table_cell(t2.rows[0].cells[i], h, bg=BG_H, bold=True, color=TITLE)

data2 = [
    ("目标", "积累1000万/1亿", "被动收入覆盖支出"),
    ("路径", "拼命工作+高位投资", "降低分母+稳定现金流"),
    ("核心指标", "账户余额", "收入/支出比率"),
    ("自由度", "等我赚够了再说", "今天就可以开始选择"),
]
for ri, (a, b, c) in enumerate(data2):
    bg = BG_A if ri % 2 == 0 else "FFFFFF"
    colored_table_cell(t2.rows[ri+1].cells[0], a, bg=BG_H, bold=True)
    colored_table_cell(t2.rows[ri+1].cells[1], b, bg=WARN_BG)
    colored_table_cell(t2.rows[ri+1].cells[2], c, bg=GREEN_BG)

spacer(doc, 6, 0)

# ===== 二、FIRE运动实践框架 =====
doc.add_page_break()
heading_para(doc, "二、FIRE运动实践框架", 1)
heading_para(doc, "▌ 核心公式", 2)
spacer(doc, 4, 2)
box_para(doc, "所需本金 = 年度开支 ÷ 4%", size=17, bold=True, color=TITLE, bg=BLUE_BG)
p_note = doc.add_paragraph()
p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_note.paragraph_format.space_before = Pt(2)
p_note.paragraph_format.space_after = Pt(6)
rn = p_note.add_run('（即"4%规则"：每年从本金中提取不超过4%，可维持30年以上资金不枯竭）')
rn.font.name = "宋体"
rn._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
rn.font.size = Pt(10)
rn.font.color.rgb = SUB

# FIRE计算表
t3 = doc.add_table(rows=5, cols=3)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
t3.style = 'Table Grid'
for i, h in enumerate(["要素", "计算方式", "示例数据"]):
    colored_table_cell(t3.rows[0].cells[i], h, bg=BG_H, bold=True, color=TITLE)
data3 = [
    ("安全提取率", "年度开支 × 25", "月支出1万 → 需300万本金"),
    ("4%规则", "1994年Bengen提出", "300万 → 年提取12万"),
    ("成功概率", "Morningstar验证", "超90%概率维持30年"),
    ("适用前提", "全球化资产配置", "指数基金+债券+黄金等分散"),
]
for ri, row in enumerate(data3):
    bg = BG_A if ri % 2 == 0 else "FFFFFF"
    for ci, val in enumerate(row):
        colored_table_cell(t3.rows[ri+1].cells[ci], val, bg=bg)

spacer(doc, 6, 4)
note_para(doc, "⚠️ 重要前提：4%规则基于美国市场历史数据，实际成功概率受市场周期、通胀水平、寿命预期等因素影响，需定期复盘调整。", color=GOLD, bg=YELLOW_BG)

# ===== 三、财富积累的数学逻辑 =====
doc.add_page_break()
heading_para(doc, "三、财富积累的数学逻辑", 1)
heading_para(doc, "▌ 复利模型（标普500历史数据）", 2)
spacer(doc, 3, 2)

t4 = doc.add_table(rows=8, cols=2)
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
t4.style = 'Table Grid'
for i, h in enumerate(["参数", "数值"]):
    colored_table_cell(t4.rows[0].cells[i], h, bg=BG_H, bold=True, color=TITLE)
data4 = [
    ("每月定投", "3,000 元"),
    ("投资周期", "25 年"),
    ("年化收益率", "10%"),
    ("累计本金", "90 万元（3,000 × 12 × 25）"),
    ("复利收益", "308 万元"),
    ("终值合计", "398 万元"),
    ("收益/本金倍数", "3.4 倍（复利威力）"),
]
for ri, (a, b) in enumerate(data4):
    bg = BG_A if ri % 2 == 0 else "FFFFFF"
    bold_row = ri in (5, 6)
    bg_row = BG_H if ri == 5 else bg
    col_b = GREEN if ri in (4, 6) else (TITLE if ri == 5 else TEXT)
    colored_table_cell(t4.rows[ri+1].cells[0], a, bg=bg_row, bold=bold_row)
    colored_table_cell(t4.rows[ri+1].cells[1], b, bg=bg_row, bold=bold_row, color=col_b)

spacer(doc, 6, 4)
heading_para(doc, "▌ 地理套利：降低生活成本", 2)
p_geo = normal_para(doc, "地理套利（Geographic Arbitrage）：利用不同城市间的生活成本差异，放大被动收入购买力。", size=22)
spacer(doc, 3, 3)

t5 = doc.add_table(rows=4, cols=3)
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
t5.style = 'Table Grid'
for i, h in enumerate(["城市类型", "月均生活支出", "所需本金（年支×25）"]):
    colored_table_cell(t5.rows[0].cells[i], h, bg=BG_H, bold=True, color=TITLE)
data5 = [
    ("一线城市（月入1万+）", "~10,000 元", "300 万", WARN_BG),
    ("二三线城市", "~5,000 元", "150 万", "FFFFFF"),
    ("大理 / 清迈 / 东南亚", "~3,000 元", "90 万", GREEN_BG),
]
for ri, row in enumerate(data5):
    a, b, c, bg = row
    colored_table_cell(t5.rows[ri+1].cells[0], a, bg=bg)
    colored_table_cell(t5.rows[ri+1].cells[1], b, bg=bg)
    colored_table_cell(t5.rows[ri+1].cells[2], c, bg=bg, bold=(ri==2), color=(GREEN if ri==2 else TEXT))

# ===== 四、生活方式设计 =====
doc.add_page_break()
heading_para(doc, "四、生活方式设计", 1)
heading_para(doc, "▌ 成本优化案例", 2)
bullet_para(doc, "北上广深 → 大理/清迈：生活成本降低 50% 以上，同时幸福感提升")
bullet_para(doc, "极简生活：重新定义\u201c足够好\u201d的标准，减少物质焦虑")
bullet_para(doc, "数字游民（Digital Nomad）：远程工作+低成本城市，效率与自由兼得")
spacer(doc, 5, 3)
heading_para(doc, "▌ 时间价值公式", 2)
spacer(doc, 3, 2)
box_para(doc, "30–60岁 黄金30年  >  延迟满足的退休生活", size=15, bold=True, color=TITLE, bg=BLUE_BG)
p_time = normal_para(doc, '真正的FIRE不是"拼命存钱等退休"，而是在黄金年龄段就拥有时间自主权。',
                     size=20, color=SUB, align=WD_ALIGN_PARAGRAPH.CENTER)

# ===== 五、执行策略 =====
doc.add_page_break()
heading_para(doc, "五、执行策略", 1)
numbered_para(doc, 1, "启动条件：每月3,000元定投全球指数基金（分散配置，降低单一市场风险）")
spacer(doc, 3, 2)
numbered_para(doc, 2, "持续要素：保持10%年化收益预期，拒绝短期市场波动干扰，坚持长期主义")
spacer(doc, 3, 2)
numbered_para(doc, 3, "终极目标：账户余额 ≠ 成功标准，可自由选择每日行程才是核心指标")
spacer(doc, 6, 4)
heading_para(doc, "▌ 三大常见陷阱", 2)

t6 = doc.add_table(rows=4, cols=2)
t6.alignment = WD_TABLE_ALIGNMENT.CENTER
t6.style = 'Table Grid'
colored_table_cell(t6.rows[0].cells[0], "陷阱", bg=BG_H, bold=True, color=TITLE)
colored_table_cell(t6.rows[0].cells[1], "说明", bg=BG_H, bold=True, color=TITLE)
traps = [
    ("高收入高消费陷阱", "收入涨了，消费也跟着涨，永远存不下本金"),
    ("投资过度集中", "ALL IN 单一个股或单一市场，黑天鹅事件导致归零"),
    ("低估生活成本", "医疗、子女、意外支出未纳入FIRE计算基础"),
]
for ri, (a, b) in enumerate(traps):
    colored_table_cell(t6.rows[ri+1].cells[0], a, bg=WARN_BG, bold=True, color=WARN)
    colored_table_cell(t6.rows[ri+1].cells[1], b, bg=WARN_BG)

# ===== 结语 =====
doc.add_page_break()
heading_para(doc, "结语：建立系统，而非追求数字", 1)
spacer(doc, 4, 4)
box_para(doc, '\u201c收入 \u2212 支出 > 0\u201d的可持续系统\n，而非追求账户里的绝对数字', size=14, bold=True, color=TITLE, bg=BLUE_BG)
spacer(doc, 6, 4)
note_para(doc, "⚠️ 数据警示：所有案例数据均为假设（如25年398万），实际收益需根据市场调整。本文档仅供思维启发，不构成投资建议。投资有风险，决策需谨慎。", color=GOLD, bg=YELLOW2_BG)
spacer(doc, 6, 4)
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_footer.paragraph_format.space_before = Pt(3)
r_f = p_footer.add_run("整理日期：2026年4月22日  |  整理工具：QClaw AI")
r_f.font.name = "宋体"
r_f._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
r_f.font.size = Pt(9)
r_f.font.color.rgb = SUB

# 保存
out_path = "/Users/lijing/.qclaw/workspace/FIRE财务自由实践指南_20260422.docx"
doc.save(out_path)
print(f"✅ 文档生成成功：{out_path}")

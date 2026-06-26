#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
保留原格式，精准替换空白占位符
- 护照号：ER606835
- 单位地址：海南省海口市秀英区秀英街道港澳大道7号
"""

import copy
import shutil
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.oxml.ns import qn

# 先复制转换后的docx（保留完整格式）
SRC = '/tmp/在职证明_converted.docx'
DST = '/Users/fasimac/Downloads/在职证明_已填写.docx'
shutil.copy2(SRC, DST)

doc = Document(DST)

# ===== 辅助函数 =====
def clone_run(para, src_run, text, bold=None):
    """复制一个run，修改文字，保留格式"""
    new_run = para.add_run(text)
    # 复制字体属性
    new_run.font.name = src_run.font.name
    new_run.font.size = src_run.font.size
    new_run.font.bold = src_run.font.bold if bold is None else bold
    new_run.font.italic = src_run.font.italic
    new_run.font.underline = src_run.font.underline
    # 复制颜色（忽略None）
    try:
        if src_run.font.color.rgb is not None:
            new_run.font.color.rgb = copy.copy(src_run.font.color.rgb)
    except Exception:
        pass
    # 设置东亚字体（保留原西文字体名）
    rPr = new_run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    orig_font = src_run.font.name
    rFonts.set(qn("w:ascii"),   orig_font if orig_font else "Times New Roman")
    rFonts.set(qn("w:hAnsi"),   orig_font if orig_font else "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "微软雅黑")
    rFonts.set(qn("w:cs"),      "微软雅黑")
    return new_run

def clear_runs(para):
    """清空所有run的文字（保留run结构）"""
    for run in para.runs:
        run.text = ""

# ============================================================
# 段落4：致送机构
# 原文本: ______驻 ______大使馆/总领事馆签证处：
# 替换为: 日本驻广州大使馆/总领事馆签证处：
# ============================================================
para4 = doc.paragraphs[4]
clear_runs(para4)
r4 = para4.runs[0]
r4.text = "日本驻广州大使馆/总领事馆签证处："
# 保留原有格式，替换为微软雅黑
r4.font.name = "微软雅黑"
rPr4 = r4._r.get_or_add_rPr()
rFonts4 = rPr4.get_or_add_rFonts()
rFonts4.set(qn("w:ascii"), "微软雅黑")
rFonts4.set(qn("w:hAnsi"), "微软雅黑")
rFonts4.set(qn("w:eastAsia"), "微软雅黑")
rFonts4.set(qn("w:cs"), "微软雅黑")

# ============================================================
# 段落5：正文长段落（完整替换）
# 原文本: ______（护照：______，出生日期……）
# ============================================================
para5 = doc.paragraphs[5]
r5 = para5.runs[0]
src_size = r5.font.size  # 152400
src_name = r5.font.name  # Times

# 整段替换文字
full_text = (
    "符发（护照：ER606835，出生日期1979年11月27日），"
    "自2016年10月迄今在我公司工作，现任职务董事长，"
    "年薪为25万元,我公司同意其自2025年01月12日至2025年01月22日"
    "前往日本旅游出行，费用由其自行承担。"
    "我司保证其在此期间遵守当地法律，并保留其职务至归国。"
)

r5.text = full_text
r5.font.name = "微软雅黑"
rPr5 = r5._r.get_or_add_rPr()
rFonts5 = rPr5.get_or_add_rFonts()
rFonts5.set(qn("w:ascii"), "微软雅黑")
rFonts5.set(qn("w:hAnsi"), "微软雅黑")
rFonts5.set(qn("w:eastAsia"), "微软雅黑")
rFonts5.set(qn("w:cs"), "微软雅黑")

# ============================================================
# 段落6：单位名称（保留标签，追加内容）
# ============================================================
para6 = doc.paragraphs[6]
clear_runs(para6)
clone_run(para6, para6.runs[0], "单位名称: ")
clone_run(para6, para6.runs[0], "华检联（海南）检测技术有限公司", bold=True)

# ============================================================
# 段落8：单位地址（保留标签，追加内容，已更新）
# ============================================================
para8 = doc.paragraphs[8]
clear_runs(para8)
clone_run(para8, para8.runs[0], "单位地址: ")
clone_run(para8, para8.runs[0], "海南省海口市秀英区秀英街道港澳大道7号", bold=True)

# ============================================================
# 段落9：领导签字
# ============================================================
para9 = doc.paragraphs[9]
clear_runs(para9)
clone_run(para9, para9.runs[0], "领导签字：")
clone_run(para9, para9.runs[0], "郑家云", bold=True)

# ============================================================
# 段落10：联系电话
# ============================================================
para10 = doc.paragraphs[10]
# runs[0]="联系电话", runs[1]="："(空)
clear_runs(para10)
clone_run(para10, para10.runs[0], "联系电话：")
clone_run(para10, para10.runs[0], "17384653707", bold=True)

# ============================================================
# 段落11：日期（保留左侧空格和"日期："）
# ============================================================
para11 = doc.paragraphs[11]
# runs[0]=空格, runs[1]=大量空格+"日期", runs[2]="：", runs[3]=空
clear_runs(para11)
# 保留原段落对齐和格式，只填入日期
# 先复制原有格式的空格
space_run = clone_run(para11, para11.runs[0], "")
date_run = clone_run(para11, para11.runs[1], "日期：")

# ============================================================
# 段落0：公司名称 → 改为微软雅黑
# ============================================================
para0 = doc.paragraphs[0]
para0.runs[0].font.name = "微软雅黑"
rPr0 = para0.runs[0]._r.get_or_add_rPr()
rFonts0 = rPr0.get_or_add_rFonts()
rFonts0.set(qn("w:ascii"), "微软雅黑")
rFonts0.set(qn("w:hAnsi"), "微软雅黑")
rFonts0.set(qn("w:eastAsia"), "微软雅黑")
rFonts0.set(qn("w:cs"), "微软雅黑")

# ============================================================
# 段落3：标题 → 改为微软雅黑
# ============================================================
para3 = doc.paragraphs[3]
para3.runs[0].font.name = "微软雅黑"
rPr3 = para3.runs[0]._r.get_or_add_rPr()
rFonts3 = rPr3.get_or_add_rFonts()
rFonts3.set(qn("w:ascii"), "微软雅黑")
rFonts3.set(qn("w:hAnsi"), "微软雅黑")
rFonts3.set(qn("w:eastAsia"), "微软雅黑")
rFonts3.set(qn("w:cs"), "微软雅黑")

# ============================================================
# 段落7：（公司公章）→ 改为微软雅黑
# ============================================================
para7 = doc.paragraphs[7]
if para7.runs:
    para7.runs[0].font.name = "微软雅黑"
    rPr7 = para7.runs[0]._r.get_or_add_rPr()
    rFonts7 = rPr7.get_or_add_rFonts()
    rFonts7.set(qn("w:eastAsia"), "微软雅黑")
    rFonts7.set(qn("w:cs"), "微软雅黑")

doc.save(DST)
print(f"文件已保存: {DST}")

# 验证
doc2 = Document(DST)
print("\n验证内容：")
for i, para in enumerate(doc2.paragraphs):
    if para.text.strip():
        print(f"  段落{i}: {para.text[:80]}")

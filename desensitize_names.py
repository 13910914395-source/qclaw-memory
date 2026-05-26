#!/usr/bin/env python3
"""
脱敏处理：
1. 金额数字 → ***
2. "华检联" → HJL
3. 姓名 → 拼音首字母
"""
import os
import re
import glob
import zipfile
import subprocess

# 配置
OUTPUT_DIR = '/Users/fasimac/.qclaw/workspace/HJL管理制度_脱敏版 5'
ZIP_PATH = '/Users/fasimac/.qclaw/workspace/HJL管理制度_脱敏版.zip'

# 姓名映射
NAME_MAP = {
    '邱雪云': 'QXY',
    '赵晶': 'ZJ',
    '符发': 'FF',
}

def name_to_initial(text):
    """替换姓名"""
    result = text
    for name, initial in NAME_MAP.items():
        # 匹配各种格式：邱雪云、邱雪云  、 邱雪云
        result = re.sub(rf'\s*{re.escape(name)}\s*', f' {initial}', result)
    return result

def desensitize_text(text):
    """脱敏处理"""
    if not text:
        return text
    
    result = text
    
    # 1. 姓名 → 首字母
    for name, initial in NAME_MAP.items():
        result = re.sub(rf'\s*{re.escape(name)}\s*', f' {initial}', result)
    
    # 2. 公司名称
    result = result.replace('华检联', 'HJL')
    
    # 3. 金额数字（增强版）
    # 标准货币格式
    result = re.sub(r'¥\s*(\d+(?:\.\d+)?)', '***', result)
    result = re.sub(r'(\d+(?:\.\d+)?)\s*元(?:/月|/年|/天|/人|/次|/小时|/平米)?', '***', result)
    result = re.sub(r'(\d+(?:\.\d+)?)\s*万\s*元', '***', result)
    result = re.sub(r'(\d+(?:\.\d+)?)\s*万', '***', result)
    # 百分比
    result = re.sub(r'(\d+(?:\.\d+)?)\s*%', '***', result)
    # 含"约"的金额
    result = re.sub(r'约\s*(\d+(?:\.\d+)?)', '***', result)
    # 独立3位以上数字
    result = re.sub(r'^\s*(\d{3,6}(?:\.\d{1,2})?)\s*$', '***', result)
    
    return result

def process_docx(src_path):
    """处理docx文件"""
    from docx import Document
    from docx.oxml.ns import qn
    
    doc = Document(src_path)
    modified = False
    
    def process_runs(paragraph):
        nonlocal modified
        for run in paragraph.runs:
            original = run.text
            new_text = desensitize_text(original)
            if new_text != original:
                run.text = new_text
                modified = True
                # 设置中文字体
                try:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                except:
                    pass
    
    # 处理所有段落
    for paragraph in doc.paragraphs:
        process_runs(paragraph)
    
    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_runs(paragraph)
    
    if modified:
        doc.save(src_path)
    
    return modified

def main():
    print('='*60)
    print('姓名脱敏处理')
    print('='*60)
    
    # 获取所有docx文件
    files = glob.glob(f'{OUTPUT_DIR}/**/*.docx', recursive=True)
    print(f'处理 {len(files)} 个文件\n')
    
    for f in sorted(files):
        filename = os.path.basename(f)
        result = process_docx(f)
        status = '✓' if result else '○'
        print(f'{status} {filename}')
    
    print('\n' + '='*60)
    print('重新打包')
    print('='*60)
    
    # 重新打包
    if os.path.exists(ZIP_PATH):
        os.remove(ZIP_PATH)
    
    file_count = 0
    with zipfile.ZipFile(ZIP_PATH, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(OUTPUT_DIR):
            dirs[:] = [d for d in dirs if not d.startswith('.')]
            for file in files:
                if file.endswith('.docx'):
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, OUTPUT_DIR)
                    zipf.write(file_path, arcname)
                    file_count += 1
    
    print(f'✓ 已更新: {ZIP_PATH}')
    print(f'  文件数: {file_count} 个')
    print(f'  大小: {os.path.getsize(ZIP_PATH) / 1024:.1f} KB')
    
    # 验证
    print('\n验证:')
    doc = Document(files[0])
    sample = doc.paragraphs[0].text if doc.paragraphs else ''
    print(f'  样本文本: {sample[:50]}...' if len(sample) > 50 else f'  样本文本: {sample}')

if __name__ == '__main__':
    main()